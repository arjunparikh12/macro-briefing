"""
AI Macro Trader — cloud web app.
Deploy on Railway. Accessible from any device, anywhere.
"""

import json
import os
import re
import uuid
import hashlib
import queue
import threading
from datetime import date, datetime
from pathlib import Path
from functools import wraps

from flask import (
    Flask, Response, jsonify, render_template_string,
    request, session, redirect, url_for
)
from werkzeug.utils import secure_filename
from apscheduler.schedulers.background import BackgroundScheduler

import data_access as db
import briefing as briefing_module
from macro_llm import get_macro_llm

app = Flask(__name__)
app.secret_key = os.environ.get("SECRET_KEY", os.urandom(24))
app.config["MAX_CONTENT_LENGTH"] = 10 * 1024 * 1024  # 10 MB max upload

# ── Paths (from shared data_access) ───────────────────────────────────────────
# All path constants and mkdir calls are in data_access.py.
# Local aliases for the few places that need direct path access:
BRIEFINGS_DIR = db.BRIEFINGS_DIR
KNOWLEDGE_DIR = db.KNOWLEDGE_DIR

# ── Auth ───────────────────────────────────────────────────────────────────────
ADMIN_PASSWORD = os.environ.get("APP_PASSWORD", "changeme")
ADMIN_USERNAME = "arjun"

def hash_password(pw: str) -> str:
    return hashlib.sha256(pw.encode()).hexdigest()

def load_users() -> dict:
    return db.load_users()

def save_users(data: dict):
    db.save_users(data)

def load_invites() -> dict:
    return db.load_invites()

def save_invites(data: dict):
    db.save_invites(data)

def login_required(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        if not session.get("logged_in"):
            if request.path.startswith("/api/"):
                return jsonify({"error": "Unauthorized"}), 401
            return redirect(url_for("login_page"))
        return f(*args, **kwargs)
    return decorated

def admin_required(f):
    """Decorator — only Arjun (admin) can access this route."""
    @wraps(f)
    def decorated(*args, **kwargs):
        if not session.get("logged_in"):
            if request.path.startswith("/api/"):
                return jsonify({"error": "Unauthorized"}), 401
            return redirect(url_for("login_page"))
        if not session.get("is_admin"):
            return jsonify({"error": "Admin only"}), 403
        return f(*args, **kwargs)
    return decorated

@app.route("/health")
def health():
    return "ok", 200

@app.route("/login", methods=["GET", "POST"])
def login_page():
    error = ""
    if request.method == "POST":
        username = request.form.get("username", "").strip().lower()
        pw = request.form.get("password", "")
        # Admin login
        if username == ADMIN_USERNAME and pw == ADMIN_PASSWORD:
            session["logged_in"] = True
            session["username"] = ADMIN_USERNAME
            session["is_admin"] = True
            return redirect(url_for("index"))
        # Regular user login
        users = load_users()
        if username in users:
            stored = users[username]
            if stored.get("password_hash") == hash_password(pw) and stored.get("active", True):
                session["logged_in"] = True
                session["username"] = username
                session["is_admin"] = False
                return redirect(url_for("index"))
        error = "Incorrect username or password."
    return render_template_string(LOGIN_HTML, error=error)

@app.route("/register", methods=["GET", "POST"])
def register_page():
    error = ""
    success = ""
    token = request.args.get("token", "") or request.form.get("token", "")
    if request.method == "POST":
        invites = load_invites()
        token = request.form.get("token", "").strip()
        username = request.form.get("username", "").strip().lower()
        pw = request.form.get("password", "")
        pw2 = request.form.get("password2", "")
        # Validate token
        if token not in invites:
            error = "Invalid or expired invite token."
        elif invites[token].get("used"):
            error = "This invite has already been used."
        elif not re.match(r"^[a-z0-9_]{3,20}$", username):
            error = "Username must be 3-20 characters: letters, numbers, underscore only."
        elif username == ADMIN_USERNAME:
            error = "That username is reserved."
        elif pw != pw2:
            error = "Passwords do not match."
        elif len(pw) < 8:
            error = "Password must be at least 8 characters."
        else:
            users = load_users()
            if username in users:
                error = "Username already taken."
            else:
                users[username] = {
                    "password_hash": hash_password(pw),
                    "created_at": datetime.now().strftime("%Y-%m-%d %H:%M"),
                    "active": True,
                    "invited_by": invites[token].get("created_by", "arjun"),
                    "note": invites[token].get("note", ""),
                }
                save_users(users)
                invites[token]["used"] = True
                invites[token]["used_by"] = username
                invites[token]["used_at"] = datetime.now().strftime("%Y-%m-%d %H:%M")
                save_invites(invites)
                success = f"Account created! You can now sign in as {username}."
    return render_template_string(REGISTER_HTML, error=error, success=success, token=token)

@app.route("/logout")
def logout():
    session.clear()
    return redirect(url_for("login_page"))

# ── User management API (admin only) ──────────────────────────────────────────
@app.route("/api/admin/invite", methods=["POST"])
@admin_required
def api_create_invite():
    note = (request.json or {}).get("note", "")
    token = str(uuid.uuid4())[:12]
    invites = load_invites()
    invites[token] = {
        "token": token,
        "note": str(note)[:100],
        "created_at": datetime.now().strftime("%Y-%m-%d %H:%M"),
        "created_by": ADMIN_USERNAME,
        "used": False,
    }
    save_invites(invites)
    base_url = request.host_url.rstrip("/")
    return jsonify({"ok": True, "token": token,
                    "link": f"{base_url}/register?token={token}"})

@app.route("/api/admin/invites", methods=["GET"])
@admin_required
def api_list_invites():
    invites = load_invites()
    return jsonify(list(invites.values()))

@app.route("/api/admin/invite/<token>/delete", methods=["POST"])
@admin_required
def api_delete_invite(token):
    invites = load_invites()
    invites.pop(token, None)
    save_invites(invites)
    return jsonify({"ok": True})

@app.route("/api/admin/users", methods=["GET"])
@admin_required
def api_list_users():
    users = load_users()
    return jsonify([
        {"username": u, "created_at": d.get("created_at",""),
         "active": d.get("active", True), "note": d.get("note","")}
        for u, d in users.items()
    ])

@app.route("/api/admin/users/<username>/toggle", methods=["POST"])
@admin_required
def api_toggle_user(username):
    if username == ADMIN_USERNAME:
        return jsonify({"error": "Cannot disable admin"}), 400
    users = load_users()
    if username not in users:
        return jsonify({"error": "Not found"}), 404
    users[username]["active"] = not users[username].get("active", True)
    save_users(users)
    return jsonify({"ok": True, "active": users[username]["active"]})

@app.route("/api/admin/users/<username>/delete", methods=["POST"])
@admin_required
def api_delete_user(username):
    if username == ADMIN_USERNAME:
        return jsonify({"error": "Cannot delete admin"}), 400
    users = load_users()
    users.pop(username, None)
    save_users(users)
    return jsonify({"ok": True})

# ── Session info ───────────────────────────────────────────────────────────────
@app.route("/api/me")
@login_required
def api_me():
    return jsonify({
        "username": session.get("username", ""),
        "is_admin": session.get("is_admin", False),
    })

# ── Helpers ────────────────────────────────────────────────────────────────────
def get_status():
    today = date.today().strftime("%Y-%m-%d")
    briefings = db.list_briefings()
    return {
        "paused": db.is_paused(),
        "today_done": db.briefing_exists(today),
        "today_date": today,
        "briefing_count": len(briefings),
        "latest_briefing": briefings[0] if briefings else None,
    }

def load_feedback():
    return db.load_feedback()

def save_feedback_data(data):
    db.save_feedback(data)

def list_briefings():
    return db.list_briefings()

# ── Generation (streaming via SSE) ────────────────────────────────────────────
_gen_lock = threading.Lock()  # prevent concurrent generation (API cost + file corruption)

@app.route("/api/generate", methods=["GET"])
@admin_required
def api_generate():
    if db.is_paused():
        def paused():
            yield "data: Briefings are paused. Resume in Settings.\n\n"
            yield "data: __DONE_ERROR__\n\n"
        return Response(paused(), mimetype="text/event-stream",
                        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"})

    if not _gen_lock.acquire(blocking=False):
        def already_running():
            yield "data: A briefing is already being generated. Please wait.\n\n"
            yield "data: __DONE_ERROR__\n\n"
        return Response(already_running(), mimetype="text/event-stream",
                        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"})

    today = date.today().strftime("%Y-%m-%d")
    output_file = BRIEFINGS_DIR / f"macro-briefing-{today}.md"

    q = queue.Queue()

    def run():
        try:
            full_text = briefing_module.generate_briefing(
                stream_callback=lambda chunk: q.put(("chunk", chunk))
            )
            output_file.write_text(full_text)
            q.put(("done", output_file.name))
        except Exception as e:
            q.put(("error", str(e)))
        finally:
            _gen_lock.release()

    t = threading.Thread(target=run, daemon=True)
    t.start()

    def stream():
        buffer = ""
        while True:
            try:
                kind, payload = q.get(timeout=120)
            except queue.Empty:
                yield "data: __DONE_ERROR__\n\n"
                return

            if kind == "chunk":
                buffer += payload
                while "\n" in buffer:
                    line, buffer = buffer.split("\n", 1)
                    yield f"data: {line}\n\n"
            elif kind == "done":
                if buffer.strip():
                    yield f"data: {buffer}\n\n"
                yield f"data: __DONE_OK__{payload}\n\n"
                return
            elif kind == "error":
                yield f"data: Error: {payload}\n\n"
                yield "data: __DONE_ERROR__\n\n"
                return

    return Response(stream(), mimetype="text/event-stream",
                    headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"})

# ── Standard routes ────────────────────────────────────────────────────────────
@app.route("/")
@login_required
def index():
    return render_template_string(HTML)

@app.route("/api/status")
@login_required
def api_status():
    return jsonify(get_status())

@app.route("/api/briefings")
@login_required
def api_briefings():
    return jsonify(list_briefings())

@app.route("/api/briefing/<filename>")
@login_required
def api_briefing(filename):
    if not re.match(r"^macro-briefing-\d{4}-\d{2}-\d{2}\.md$", filename):
        return jsonify({"error": "Invalid filename"}), 400
    path = BRIEFINGS_DIR / filename
    if not path.exists():
        return jsonify({"error": "Not found"}), 404
    return jsonify({"content": path.read_text()})

@app.route("/api/delete-today", methods=["POST"])
@admin_required
def api_delete_today():
    today = date.today().strftime("%Y-%m-%d")
    path = BRIEFINGS_DIR / f"macro-briefing-{today}.md"
    if path.exists():
        path.unlink()
    return jsonify({"ok": True})

@app.route("/api/pause", methods=["POST"])
@admin_required
def api_pause():
    db.set_paused(True)
    return jsonify({"ok": True})

@app.route("/api/resume", methods=["POST"])
@admin_required
def api_resume():
    db.set_paused(False)
    return jsonify({"ok": True})

@app.route("/api/feedback", methods=["GET"])
@login_required
def api_get_feedback():
    filename = request.args.get("filename", "")
    if not re.match(r"^macro-briefing-\d{4}-\d{2}-\d{2}\.md$", filename):
        return jsonify({"error": "Invalid"}), 400
    date_key = filename.replace("macro-briefing-", "").replace(".md", "")
    return jsonify(load_feedback().get(date_key, []))

@app.route("/api/feedback", methods=["POST"])
@admin_required
def api_save_feedback():
    body = request.json
    filename = body.get("filename", "")
    if not re.match(r"^macro-briefing-\d{4}-\d{2}-\d{2}\.md$", filename):
        return jsonify({"error": "Invalid"}), 400
    date_key = filename.replace("macro-briefing-", "").replace(".md", "")
    entries = []
    for e in body.get("entries", []):
        rating = e.get("rating") if e.get("rating") in ("up", "down", None) else None
        entry = {
            "rating": rating,
            "note": str(e.get("note", ""))[:1000],
        }
        if e.get("section"):
            entry["section"] = str(e.get("section", ""))[:200]
        else:
            entry["trade"] = str(e.get("trade", ""))[:500]
        entries.append(entry)
    data = load_feedback()
    existing = data.get(date_key, [])
    new_sections = {e["section"]: e for e in entries if e.get("section")}
    new_trades = [e for e in entries if not e.get("section")]
    kept = [e for e in existing if e.get("section") and e["section"] not in new_sections]
    if new_trades:
        data[date_key] = kept + list(new_sections.values()) + new_trades
    else:
        existing_trades = [e for e in existing if not e.get("section")]
        data[date_key] = kept + list(new_sections.values()) + existing_trades
    save_feedback_data(data)
    return jsonify({"ok": True})

# ── Document knowledge base ────────────────────────────────────────────────────
ALLOWED_EXTENSIONS = {".pdf", ".docx", ".txt"}

def extract_text_from_file(filepath: Path) -> str:
    suffix = filepath.suffix.lower()
    try:
        if suffix == ".pdf":
            import pdfplumber
            pages = []
            with pdfplumber.open(filepath) as pdf:
                for page in pdf.pages[:25]:
                    text = page.extract_text()
                    if text:
                        pages.append(text)
            return "\n".join(pages)
        elif suffix == ".docx":
            from docx import Document
            doc = Document(filepath)
            return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        else:
            return filepath.read_text(errors="replace")[:60000]
    except Exception as e:
        return f"[Text extraction error: {e}]"

DOC_TYPES = {"tactical", "guide", "reference"}

def summarize_document(title: str, raw_text: str) -> str:
    """Summarize a document for knowledge base injection."""
    from anthropic import Anthropic
    haiku_client = Anthropic(api_key=os.environ["ANTHROPIC_API_KEY"])
    resp = haiku_client.messages.create(
        model="claude-haiku-4-5",
        max_tokens=800,
        messages=[{
            "role": "user",
            "content": (
                "You are summarizing a document for use in a daily macro briefing system "
                "used by a QIS structurer focused on rates, FX, and cross-currency basis.\n\n"
                "Output a concise structured note (300-500 words) covering:\n"
                "- Core thesis, macro narrative, or analytical framework\n"
                "- Key variables and signals to monitor\n"
                "- Analytical frameworks or reasoning patterns described\n"
                "- Any quantitative rules, thresholds, or z-score logic\n"
                "- How the analytical lens in this document should inform daily briefing preparation\n\n"
                f"Document title: {title}\n\n"
                f"Document content:\n{raw_text[:9000]}"
            )
        }]
    )
    return resp.content[0].text

def list_documents() -> list:
    docs = []
    for f in sorted(KNOWLEDGE_DIR.glob("*.json"), reverse=True):
        try:
            with open(f) as fp:
                doc = json.load(fp)
            docs.append({
                "id": doc["id"],
                "title": doc["title"],
                "filename": doc["filename"],
                "uploaded_at": doc["uploaded_at"],
                "active": doc.get("active", True),
                "doc_type": doc.get("doc_type", "guide"),
            })
        except Exception:
            continue
    return docs

@app.route("/api/upload", methods=["POST"])
@admin_required
def api_upload():
    if "file" not in request.files:
        return jsonify({"error": "No file provided"}), 400
    f = request.files["file"]
    if not f.filename:
        return jsonify({"error": "No filename"}), 400
    suffix = Path(f.filename).suffix.lower()
    if suffix not in ALLOWED_EXTENSIONS:
        return jsonify({"error": f"File type {suffix} not allowed. Use .pdf, .docx, or .txt"}), 400
    safe_name = secure_filename(f.filename)
    doc_id = str(uuid.uuid4())[:8]
    tmp_path = DATA_DIR / f"tmp_{doc_id}{suffix}"
    try:
        KNOWLEDGE_DIR.mkdir(parents=True, exist_ok=True)
        f.save(str(tmp_path))
        raw_text = extract_text_from_file(tmp_path)
        if not raw_text.strip():
            return jsonify({"error": "Could not extract text from file"}), 422
        title = Path(f.filename).stem.replace("_", " ").replace("-", " ")
        doc_type = request.form.get("doc_type", "guide")
        if doc_type not in DOC_TYPES:
            doc_type = "guide"
        print(f"[upload] Extracted {len(raw_text)} chars from '{title}', type={doc_type}, calling Haiku...")
        summary = summarize_document(title, raw_text)
        print(f"[upload] Haiku summary done ({len(summary)} chars)")
        doc_record = {
            "id": doc_id,
            "title": title,
            "filename": safe_name,
            "summary": summary,
            "doc_type": doc_type,
            "uploaded_at": datetime.now().strftime("%Y-%m-%d %H:%M"),
            "active": True,
        }
        kb_file = KNOWLEDGE_DIR / f"{doc_id}.json"
        kb_file.write_text(json.dumps(doc_record, indent=2))
        return jsonify({"ok": True, "doc": {
            "id": doc_id, "title": title, "filename": safe_name,
            "uploaded_at": doc_record["uploaded_at"], "active": True,
            "doc_type": doc_type,
        }})
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500
    finally:
        if tmp_path.exists():
            tmp_path.unlink()

@app.route("/api/documents", methods=["GET"])
@login_required
def api_list_documents():
    return jsonify(list_documents())

@app.route("/api/documents/<doc_id>/toggle", methods=["POST"])
@admin_required
def api_toggle_document(doc_id):
    if not re.match(r"^[a-f0-9]{8}$", doc_id):
        return jsonify({"error": "Invalid id"}), 400
    kb_file = KNOWLEDGE_DIR / f"{doc_id}.json"
    if not kb_file.exists():
        return jsonify({"error": "Not found"}), 404
    with open(kb_file) as fp:
        doc = json.load(fp)
    doc["active"] = not doc.get("active", True)
    kb_file.write_text(json.dumps(doc, indent=2))
    return jsonify({"ok": True, "active": doc["active"]})

@app.route("/api/documents/<doc_id>/type", methods=["POST"])
@admin_required
def api_set_doc_type(doc_id):
    if not re.match(r"^[a-f0-9]{8}$", doc_id):
        return jsonify({"error": "Invalid id"}), 400
    kb_file = KNOWLEDGE_DIR / f"{doc_id}.json"
    if not kb_file.exists():
        return jsonify({"error": "Not found"}), 404
    new_type = (request.json or {}).get("doc_type", "guide")
    if new_type not in DOC_TYPES:
        return jsonify({"error": "Must be 'tactical', 'guide', or 'reference'"}), 400
    with open(kb_file) as fp:
        doc = json.load(fp)
    doc["doc_type"] = new_type
    kb_file.write_text(json.dumps(doc, indent=2))
    return jsonify({"ok": True, "doc_type": new_type})

@app.route("/api/documents/<doc_id>/delete", methods=["POST"])
@admin_required
def api_delete_document(doc_id):
    if not re.match(r"^[a-f0-9]{8}$", doc_id):
        return jsonify({"error": "Invalid id"}), 400
    kb_file = KNOWLEDGE_DIR / f"{doc_id}.json"
    if kb_file.exists():
        kb_file.unlink()
    return jsonify({"ok": True})

# ── Chat / Q&A on briefings ───────────────────────────────────────────────────

def load_chat_history(briefing_date: str) -> list:
    return db.load_chat_history(briefing_date)

def save_chat_history(briefing_date: str, messages: list):
    db.save_chat_history(briefing_date, messages)

def load_insights() -> list:
    return db.load_insights()

def save_insights(insights: list):
    db.save_insights(insights)

@app.route("/api/chat/history/<briefing_date>")
@admin_required
def api_chat_history(briefing_date):
    if not re.match(r"^\d{4}-\d{2}-\d{2}$", briefing_date):
        return jsonify({"error": "Invalid date"}), 400
    return jsonify(load_chat_history(briefing_date))

@app.route("/api/chat", methods=["POST"])
@admin_required
def api_chat():
    """Stream a chat response about a briefing section. Uses MacroLLM (no external API)."""
    import time
    body = request.json or {}
    briefing_date = body.get("briefing_date", "")
    user_message = body.get("message", "").strip()
    section_context = body.get("section_context", "")

    if not user_message or not briefing_date:
        return jsonify({"error": "Missing message or briefing_date"}), 400
    if not re.match(r"^\d{4}-\d{2}-\d{2}$", briefing_date):
        return jsonify({"error": "Invalid date"}), 400

    # Load the briefing content
    briefing_path = BRIEFINGS_DIR / f"macro-briefing-{briefing_date}.md"
    briefing_content = briefing_path.read_text() if briefing_path.exists() else ""

    # Load chat history for context
    chat_history = load_chat_history(briefing_date)

    # Generate response using MacroLLM (no Claude API call)
    llm = get_macro_llm()

    def stream():
        try:
            full_response = llm.ask(
                briefing_content=briefing_content,
                section_context=section_context,
                question=user_message,
                chat_history=chat_history,
            )

            # Stream by line with brief pauses — keeps the frontend UX
            # but doesn't block threads as long as word-by-word
            lines = full_response.split("\n")
            for i, line in enumerate(lines):
                chunk = line if i == 0 else "\n" + line
                yield f"data: {json.dumps({'chunk': chunk})}\n\n"
                time.sleep(0.02)

            # Save to chat history
            chat_history.append({"role": "user", "content": user_message, "ts": datetime.now().strftime("%H:%M")})
            chat_history.append({"role": "assistant", "content": full_response, "ts": datetime.now().strftime("%H:%M")})
            save_chat_history(briefing_date, chat_history)

            yield f"data: {json.dumps({'done': True})}\n\n"
        except Exception as e:
            import traceback
            traceback.print_exc()
            yield f"data: {json.dumps({'error': str(e)})}\n\n"

    return Response(stream(), mimetype="text/event-stream",
                    headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"})

@app.route("/api/chat/feedback", methods=["POST"])
@admin_required
def api_chat_feedback():
    """Record good/bad feedback on the last MacroLLM response."""
    body = request.json or {}
    feedback = body.get("feedback", "")
    if feedback not in ("good", "bad"):
        return jsonify({"error": "feedback must be 'good' or 'bad'"}), 400
    llm = get_macro_llm()
    llm.give_feedback(feedback)
    return jsonify({"ok": True})

@app.route("/api/chat/save-insight", methods=["POST"])
@admin_required
def api_save_insight():
    """Save a key takeaway from a chat conversation as a permanent insight."""
    body = request.json or {}
    insight_text = body.get("insight", "").strip()
    briefing_date = body.get("briefing_date", "")
    if not insight_text:
        return jsonify({"error": "No insight text"}), 400
    insights = load_insights()
    insights.append({
        "insight": insight_text,
        "date": briefing_date,
        "saved_at": datetime.now().strftime("%Y-%m-%d %H:%M"),
    })
    # Keep last 100 insights
    if len(insights) > 100:
        insights = insights[-100:]
    save_insights(insights)
    return jsonify({"ok": True})

@app.route("/api/insights", methods=["GET"])
@admin_required
def api_get_insights():
    return jsonify(load_insights())

@app.route("/api/insights/<int:idx>/delete", methods=["POST"])
@admin_required
def api_delete_insight(idx):
    insights = load_insights()
    if 0 <= idx < len(insights):
        insights.pop(idx)
        save_insights(insights)
    return jsonify({"ok": True})

# ── Scheduler (auto-generate at 6 AM ET Mon-Fri) ──────────────────────────────
def scheduled_generate():
    if db.is_paused():
        return
    today = date.today().strftime("%Y-%m-%d")
    if db.briefing_exists(today):
        return
    try:
        text = briefing_module.generate_briefing()
        db.write_briefing(today, text)
        print(f"[scheduler] Briefing generated: macro-briefing-{today}.md")
    except Exception as e:
        print(f"[scheduler] Generation failed: {e}")

try:
    scheduler = BackgroundScheduler(timezone="America/New_York")
    scheduler.add_job(scheduled_generate, "cron", day_of_week="mon-fri", hour=6, minute=0)
    scheduler.start()
except Exception as e:
    print(f"[scheduler] Failed to start: {e}")

# ── Login page HTML ────────────────────────────────────────────────────────────
LOGIN_HTML = r"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>AI Macro Trader — Sign In</title>
<style>
  * { box-sizing: border-box; margin: 0; padding: 0; }
  body { background: #0d0d0d; color: #e8e8e8; font-family: -apple-system, sans-serif;
         display: flex; align-items: center; justify-content: center; min-height: 100vh; }
  .box { background: #161616; border: 1px solid #2a2a2a; border-radius: 12px;
         padding: 40px 36px; width: 100%; max-width: 360px; }
  .brand { font-size: 22px; font-weight: 700; color: #c8a96e; margin-bottom: 4px; letter-spacing: -0.02em; }
  .sub   { font-size: 13px; color: #555; margin-bottom: 28px; }
  label  { display: block; font-size: 12px; color: #666; margin-bottom: 5px; }
  input  { width: 100%; background: #0d0d0d; border: 1px solid #2a2a2a;
           border-radius: 8px; padding: 11px 14px; color: #e8e8e8; font-size: 15px;
           margin-bottom: 12px; }
  input:focus { outline: none; border-color: #c8a96e; }
  button { width: 100%; background: #c8a96e; color: #000; border: none; border-radius: 8px;
           padding: 12px; font-size: 15px; font-weight: 600; cursor: pointer; margin-top: 2px; }
  .error { color: #e05252; font-size: 13px; margin-top: 12px; }
  .reg-link { text-align: center; margin-top: 18px; font-size: 12px; color: #444; }
  .reg-link a { color: #c8a96e; text-decoration: none; }
</style>
</head>
<body>
<div class="box">
  <div class="brand">AI Macro Trader</div>
  <div class="sub">Rates · FX · Cross-Currency Basis</div>
  <form method="post">
    <label>Username</label>
    <input type="text" name="username" placeholder="username" autocomplete="username" autofocus>
    <label>Password</label>
    <input type="password" name="password" placeholder="Password" autocomplete="current-password">
    <button type="submit">Sign in</button>
    {% if error %}<div class="error">{{ error }}</div>{% endif %}
  </form>
  <div class="reg-link">Have an invite? <a href="/register">Create account</a></div>
</div>
</body>
</html>"""

# ── Register page HTML ─────────────────────────────────────────────────────────
REGISTER_HTML = r"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>AI Macro Trader — Create Account</title>
<style>
  * { box-sizing: border-box; margin: 0; padding: 0; }
  body { background: #0d0d0d; color: #e8e8e8; font-family: -apple-system, sans-serif;
         display: flex; align-items: center; justify-content: center; min-height: 100vh; }
  .box { background: #161616; border: 1px solid #2a2a2a; border-radius: 12px;
         padding: 40px 36px; width: 100%; max-width: 380px; }
  .brand { font-size: 22px; font-weight: 700; color: #c8a96e; margin-bottom: 4px; letter-spacing: -0.02em; }
  .sub   { font-size: 13px; color: #555; margin-bottom: 24px; }
  label  { display: block; font-size: 12px; color: #666; margin-bottom: 5px; }
  input  { width: 100%; background: #0d0d0d; border: 1px solid #2a2a2a;
           border-radius: 8px; padding: 11px 14px; color: #e8e8e8; font-size: 15px;
           margin-bottom: 12px; }
  input:focus { outline: none; border-color: #c8a96e; }
  button { width: 100%; background: #c8a96e; color: #000; border: none; border-radius: 8px;
           padding: 12px; font-size: 15px; font-weight: 600; cursor: pointer; margin-top: 2px; }
  .error   { color: #e05252; font-size: 13px; margin-top: 12px; }
  .success { color: #4caf6e; font-size: 13px; margin-top: 12px; }
  .note    { font-size: 12px; color: #555; margin-bottom: 18px; line-height: 1.5; }
  .back    { text-align: center; margin-top: 18px; font-size: 12px; color: #444; }
  .back a  { color: #c8a96e; text-decoration: none; }
</style>
</head>
<body>
<div class="box">
  <div class="brand">AI Macro Trader</div>
  <div class="sub">Create your account</div>
  <p class="note">You need an invite token from Arjun to register.</p>
  <form method="post">
    <label>Invite Token</label>
    <input type="text" name="token" placeholder="e.g. a1b2c3d4e5f6" value="{{ token }}" autocomplete="off">
    <label>Choose a Username</label>
    <input type="text" name="username" placeholder="lowercase, no spaces" autocomplete="username">
    <label>Password</label>
    <input type="password" name="password" placeholder="At least 8 characters" autocomplete="new-password">
    <label>Confirm Password</label>
    <input type="password" name="password2" placeholder="Repeat password" autocomplete="new-password">
    <button type="submit">Create Account</button>
    {% if error %}<div class="error">{{ error }}</div>{% endif %}
    {% if success %}<div class="success">{{ success }} <a href="/login" style="color:#c8a96e">Sign in →</a></div>{% endif %}
  </form>
  <div class="back"><a href="/login">← Back to sign in</a></div>
</div>
</body>
</html>"""

# ── Main app HTML ──────────────────────────────────────────────────────────────
HTML = r"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>AI Macro Trader</title>
<style>
  :root {
    --bg: #0d0d0d; --surface: #161616; --border: #2a2a2a;
    --accent: #c8a96e; --accent-dim: #8a6e3e;
    --text: #e8e8e8; --muted: #666; --green: #4caf6e; --red: #e05252; --radius: 10px;
  }
  * { box-sizing: border-box; margin: 0; padding: 0; }
  body { background: var(--bg); color: var(--text);
         font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", sans-serif;
         font-size: 15px; min-height: 100vh; }
  .app { display: flex; flex-direction: column; min-height: 100vh; }
  header { padding: 16px 20px; border-bottom: 1px solid var(--border);
           display: flex; align-items: center; justify-content: space-between;
           background: var(--surface); position: sticky; top: 0; z-index: 10; }
  .logo { font-size: 17px; font-weight: 700; letter-spacing: -0.01em; color: var(--accent); }
  .logo span { color: var(--muted); font-weight: 400; font-size: 13px; margin-left: 6px; }
  .header-right { display: flex; gap: 10px; align-items: center; }
  .header-user { font-size: 12px; color: var(--muted); }
  main { flex: 1; padding: 20px; max-width: 900px; width: 100%; margin: 0 auto; }

  .btn { display: inline-flex; align-items: center; justify-content: center; gap: 7px;
         padding: 10px 18px; border-radius: var(--radius); font-size: 14px; font-weight: 500;
         cursor: pointer; border: none; transition: opacity 0.15s, transform 0.1s; white-space: nowrap; }
  .btn:active { transform: scale(0.97); }
  .btn:disabled { opacity: 0.4; cursor: not-allowed; transform: none; }
  .btn-primary { background: var(--accent); color: #000; }
  .btn-secondary { background: var(--surface); color: var(--text); border: 1px solid var(--border); }
  .btn-danger { background: rgba(224,82,82,0.15); color: var(--red); border: 1px solid rgba(224,82,82,0.3); }
  .btn-sm { padding: 7px 13px; font-size: 13px; }
  .btn-large { padding: 14px 28px; font-size: 16px; font-weight: 600; width: 100%; }

  .card { background: var(--surface); border: 1px solid var(--border);
          border-radius: var(--radius); padding: 20px; margin-bottom: 16px; }
  .card-title { font-size: 11px; font-weight: 600; letter-spacing: 0.08em;
                text-transform: uppercase; color: var(--muted); margin-bottom: 14px; }

  .pill { display: inline-flex; align-items: center; gap: 5px;
          padding: 3px 10px; border-radius: 20px; font-size: 12px; font-weight: 500; }
  .pill-green { background: rgba(76,175,110,0.15); color: var(--green); }
  .pill-red   { background: rgba(224,82,82,0.15);  color: var(--red); }
  .pill-gold  { background: rgba(200,169,110,0.15); color: var(--accent); }
  .dot { width: 7px; height: 7px; border-radius: 50%; background: currentColor; }

  .generate-area { text-align: center; padding: 10px 0 4px; }
  .generate-area .sub { color: var(--muted); font-size: 13px; margin-top: 8px; }

  .log-box { background: #0a0a0a; border: 1px solid var(--border); border-radius: 8px;
             padding: 14px 16px; font-family: "SF Mono", "Menlo", monospace;
             font-size: 12px; color: #aaa; max-height: 220px; overflow-y: auto;
             display: none; margin-top: 14px; line-height: 1.6; }
  .log-box.visible { display: block; }
  .log-line-ok  { color: var(--green); }
  .log-line-err { color: var(--red); }

  .briefing-content { font-family: "Georgia", serif; font-size: 15px; line-height: 1.75;
                      color: #d8d8d8; white-space: pre-wrap; word-break: break-word; }
  .briefing-content h1 { font-size: 20px; color: var(--accent); margin-bottom: 16px;
                         font-family: -apple-system, sans-serif; }
  .briefing-content h2 { font-size: 16px; color: var(--text); margin: 22px 0 8px;
                         font-family: -apple-system, sans-serif;
                         border-bottom: 1px solid var(--border); padding-bottom: 5px; }
  .briefing-content h3 { font-size: 14px; color: var(--accent-dim); margin: 16px 0 6px;
                         font-family: -apple-system, sans-serif; }

  .history-item { display: flex; align-items: center; justify-content: space-between;
                  padding: 10px 0; border-bottom: 1px solid var(--border); gap: 10px; }
  .history-item:last-child { border-bottom: none; }
  .history-date { font-size: 14px; font-weight: 500; }
  .history-actions { display: flex; gap: 8px; }

  /* Trade Feedback */
  .feedback-card { background: #0f0f0f; border: 1px solid var(--border);
                   border-radius: var(--radius); padding: 18px 20px; margin-bottom: 16px; }
  .feedback-card-title { font-size: 11px; font-weight: 600; letter-spacing: 0.08em;
                         text-transform: uppercase; color: var(--accent); margin-bottom: 14px;
                         display: flex; align-items: center; gap: 8px; }
  .trade-feedback-item { border: 1px solid var(--border); border-radius: 8px;
                         padding: 14px 16px; margin-bottom: 10px; background: var(--surface); }
  .trade-text { font-size: 13px; color: #ccc; line-height: 1.5; margin-bottom: 10px; white-space: pre-wrap; }
  .feedback-controls { display: flex; align-items: center; gap: 8px; flex-wrap: wrap; }
  .thumb-btn { background: none; border: 1px solid var(--border); border-radius: 6px;
               padding: 5px 12px; cursor: pointer; font-size: 16px; transition: all 0.15s; color: var(--muted); }
  .thumb-btn:hover { border-color: var(--accent); }
  .thumb-btn.active-up   { background: rgba(76,175,110,0.15); border-color: var(--green); color: var(--green); }
  .thumb-btn.active-down { background: rgba(224,82,82,0.15);  border-color: var(--red);   color: var(--red); }
  .feedback-note { flex: 1; min-width: 160px; background: #0d0d0d; border: 1px solid var(--border);
                   border-radius: 6px; padding: 6px 10px; color: var(--text); font-size: 12px;
                   font-family: inherit; resize: none; height: 34px; line-height: 1.4; transition: border-color 0.15s; }
  .feedback-note:focus { outline: none; border-color: var(--accent); height: 60px; }
  .feedback-save-btn { background: var(--accent); color: #000; border: none; border-radius: 6px;
                       padding: 7px 16px; font-size: 13px; font-weight: 600; cursor: pointer; }
  .feedback-save-btn:disabled { opacity: 0.4; cursor: not-allowed; }
  .feedback-saved-msg { font-size: 12px; color: var(--green); display: none; }

  /* Section Feedback (inline in briefing) */
  .section-feedback-row { display: flex; align-items: center; gap: 7px; margin: 4px 0 18px;
                          padding: 8px 12px; border-radius: 7px; background: rgba(255,255,255,0.03);
                          border: 1px solid transparent; transition: border-color 0.2s; }
  .section-feedback-row:hover { border-color: var(--border); }
  .section-feedback-row .sf-label { font-size: 11px; color: var(--muted); font-family: -apple-system, sans-serif;
                                     letter-spacing: 0.04em; white-space: nowrap; }
  .sf-thumb { background: none; border: 1px solid var(--border); border-radius: 5px;
              padding: 3px 9px; cursor: pointer; font-size: 13px; transition: all 0.15s;
              color: var(--muted); line-height: 1.4; }
  .sf-thumb:hover { border-color: var(--accent); }
  .sf-thumb.sf-up   { background: rgba(76,175,110,0.15); border-color: var(--green); color: var(--green); }
  .sf-thumb.sf-down { background: rgba(224,82,82,0.15);  border-color: var(--red);   color: var(--red); }
  .sf-note { flex: 1; min-width: 120px; max-width: 360px; background: #0d0d0d;
             border: 1px solid var(--border); border-radius: 5px; padding: 4px 8px;
             color: var(--text); font-size: 12px; font-family: inherit; resize: none;
             height: 28px; line-height: 1.4; transition: border-color 0.15s, height 0.15s; }
  .sf-note:focus { outline: none; border-color: var(--accent); height: 52px; }
  .sf-save { background: none; border: 1px solid var(--border); border-radius: 5px;
             padding: 3px 10px; cursor: pointer; font-size: 11px; color: var(--muted);
             transition: all 0.15s; white-space: nowrap; }
  .sf-save:hover { border-color: var(--accent); color: var(--accent); }
  .sf-saved { font-size: 11px; color: var(--green); display: none; }

  /* Knowledge Base / Documents */
  .doc-item { display: flex; align-items: center; gap: 10px; padding: 10px 0;
              border-bottom: 1px solid var(--border); flex-wrap: wrap; }
  .doc-item:last-child { border-bottom: none; }
  .doc-name { font-size: 13px; font-weight: 500; flex: 1; min-width: 120px; word-break: break-word; }
  .doc-date { font-size: 11px; color: var(--muted); white-space: nowrap; }
  .doc-actions { display: flex; gap: 6px; align-items: center; }
  .toggle-pill { padding: 3px 10px; border-radius: 20px; font-size: 11px; font-weight: 600;
                 cursor: pointer; border: none; transition: all 0.15s; }
  .toggle-pill.active   { background: rgba(76,175,110,0.15); color: var(--green); }
  .toggle-pill.inactive { background: rgba(200,169,110,0.1); color: var(--muted); }
  .doc-type-select { padding: 3px 6px; border-radius: 6px; font-size: 11px; font-weight: 600;
                     cursor: pointer; border: 1px solid var(--border); transition: all 0.15s;
                     background: #0f0f0f; color: var(--text); outline: none; -webkit-appearance: auto; }
  .doc-type-select:focus { border-color: var(--accent); }
  .doc-type-select.dt-tactical  { color: #ffa500; border-color: rgba(255,165,0,0.3); }
  .doc-type-select.dt-guide     { color: #6495ed; border-color: rgba(100,149,237,0.3); }
  .doc-type-select.dt-reference { color: #a78bfa; border-color: rgba(167,139,250,0.3); }
  .doc-delete { background: none; border: 1px solid var(--border); border-radius: 5px;
                padding: 3px 9px; cursor: pointer; font-size: 12px; color: var(--muted);
                transition: all 0.15s; }
  .doc-delete:hover { border-color: var(--red); color: var(--red); }
  .upload-area { margin-top: 12px; }
  .upload-status { font-size: 12px; color: var(--muted); margin-top: 8px; min-height: 18px; }
  .upload-status.ok  { color: var(--green); }
  .upload-status.err { color: var(--red); }

  /* Users / Invites */
  .user-item { display: flex; align-items: center; gap: 10px; padding: 10px 0;
               border-bottom: 1px solid var(--border); flex-wrap: wrap; }
  .user-item:last-child { border-bottom: none; }
  .user-name { font-size: 13px; font-weight: 500; flex: 1; }
  .user-date { font-size: 11px; color: var(--muted); }
  .user-note { font-size: 11px; color: var(--muted); font-style: italic; }
  .invite-item { display: flex; align-items: center; gap: 10px; padding: 10px 0;
                 border-bottom: 1px solid var(--border); flex-wrap: wrap; }
  .invite-item:last-child { border-bottom: none; }
  .invite-token { font-family: "SF Mono", monospace; font-size: 12px; color: var(--accent);
                  flex: 1; word-break: break-all; }
  .invite-meta { font-size: 11px; color: var(--muted); }
  .invite-link { font-size: 11px; color: var(--muted); word-break: break-all; }
  .copy-btn { background: none; border: 1px solid var(--border); border-radius: 5px;
              padding: 3px 9px; cursor: pointer; font-size: 11px; color: var(--muted);
              transition: all 0.15s; white-space: nowrap; }
  .copy-btn:hover { border-color: var(--accent); color: var(--accent); }
  .copy-btn.copied { border-color: var(--green); color: var(--green); }
  .invite-note-input { background: #0d0d0d; border: 1px solid var(--border); border-radius: 6px;
                       padding: 7px 10px; color: var(--text); font-size: 13px; font-family: inherit;
                       flex: 1; min-width: 140px; }
  .invite-note-input:focus { outline: none; border-color: var(--accent); }
  .status-msg { font-size: 12px; margin-top: 8px; min-height: 18px; }
  .status-msg.ok  { color: var(--green); }
  .status-msg.err { color: var(--red); }

  /* Storage steps */
  .storage-steps { font-size: 12px; color: var(--muted); line-height: 2;
                   background: #0a0a0a; border: 1px solid var(--border); border-radius: 7px;
                   padding: 12px 14px; margin-top: 10px; }
  .storage-steps strong { color: var(--text); }
  .storage-steps code { background: #1a1a1a; padding: 1px 5px; border-radius: 4px;
                        font-family: "SF Mono", monospace; font-size: 11px; color: var(--accent); }

  /* Settings form */
  .form-group { margin-bottom: 14px; }
  .form-group label { display: block; font-size: 12px; color: var(--muted); margin-bottom: 6px; }
  .form-group input { width: 100%; background: #0d0d0d; border: 1px solid var(--border);
                      border-radius: 8px; padding: 10px 12px; color: var(--text); font-size: 14px; }
  .form-group input:focus { outline: none; border-color: var(--accent); }
  .form-note { font-size: 12px; color: var(--muted); margin-top: 4px; }

  /* Tab bar */
  .tab-bar { display: none; position: fixed; bottom: 0; left: 0; right: 0;
             background: var(--surface); border-top: 1px solid var(--border);
             padding: 8px 0 env(safe-area-inset-bottom, 8px); z-index: 20; }
  .tab-bar-inner { display: flex; justify-content: space-around; }
  .tab-btn { display: flex; flex-direction: column; align-items: center; gap: 3px;
             background: none; border: none; color: var(--muted); font-size: 10px; cursor: pointer; padding: 4px 16px; }
  .tab-btn.active { color: var(--accent); }
  .tab-btn svg { width: 22px; height: 22px; }

  @media (max-width: 640px) {
    header { padding: 12px 16px; }
    main { padding: 14px 14px 80px; }
    .tab-bar { display: block; }
    .desktop-nav { display: none; }
    .btn-large { font-size: 15px; }
  }

  .section { display: none; }
  .section.active { display: block; }
  .row { display: flex; gap: 10px; align-items: center; flex-wrap: wrap; }
  .spacer { flex: 1; }
  .empty { color: var(--muted); font-size: 14px; text-align: center; padding: 30px 0; }
  .spinner { width: 16px; height: 16px; border: 2px solid rgba(200,169,110,0.3);
             border-top-color: var(--accent); border-radius: 50%;
             animation: spin 0.7s linear infinite; display: inline-block; }
  @keyframes spin { to { transform: rotate(360deg); } }
  .admin-only { display: none; }
  .admin-only.visible { display: block; }

  /* Chat panel */
  .chat-overlay { display: none; position: fixed; inset: 0; background: rgba(0,0,0,0.5); z-index: 50; }
  .chat-overlay.open { display: block; }
  .chat-panel { display: none; position: fixed; bottom: 0; left: 0; right: 0; z-index: 51;
                background: var(--surface); border-top: 2px solid var(--accent);
                max-height: 70vh; height: 70vh; flex-direction: column;
                border-radius: 16px 16px 0 0; overflow: hidden; }
  .chat-panel.open { display: flex; }
  .chat-header { display: flex; align-items: center; justify-content: space-between;
                 padding: 12px 16px; border-bottom: 1px solid var(--border); flex-shrink: 0; }
  .chat-header-title { font-size: 12px; font-weight: 600; letter-spacing: 0.05em;
                       text-transform: uppercase; color: var(--accent); }
  .chat-close { background: none; border: none; color: var(--muted); cursor: pointer;
                font-size: 18px; padding: 4px 8px; }
  .chat-close:hover { color: var(--text); }
  .chat-context { font-size: 11px; color: var(--muted); padding: 8px 16px;
                  border-bottom: 1px solid var(--border); background: #0a0a0a;
                  max-height: 60px; overflow: hidden; flex-shrink: 0; }
  .chat-messages { flex: 1; overflow-y: auto; padding: 12px 16px; }
  .chat-msg { margin-bottom: 12px; }
  .chat-msg-user { text-align: right; }
  .chat-msg-user .chat-bubble { background: rgba(200,169,110,0.15); color: var(--text);
                                display: inline-block; padding: 8px 14px; border-radius: 12px 12px 4px 12px;
                                max-width: 85%; text-align: left; font-size: 14px; line-height: 1.5; }
  .chat-msg-assistant .chat-bubble { background: #1a1a1a; color: var(--text);
                                     display: inline-block; padding: 8px 14px; border-radius: 12px 12px 12px 4px;
                                     max-width: 85%; font-size: 14px; line-height: 1.5; }
  .chat-msg-assistant .chat-bubble strong { color: var(--accent); }
  .chat-msg-time { font-size: 10px; color: var(--muted); margin-top: 2px; }
  .chat-typing { color: var(--muted); font-size: 13px; font-style: italic; padding: 4px 0; }
  .chat-input-row { display: flex; gap: 8px; padding: 12px 16px; border-top: 1px solid var(--border);
                    background: var(--surface); flex-shrink: 0; }
  .chat-input { flex: 1; background: #0d0d0d; border: 1px solid var(--border); border-radius: 8px;
                padding: 10px 14px; color: var(--text); font-size: 14px; font-family: inherit;
                resize: none; max-height: 100px; }
  .chat-input:focus { outline: none; border-color: var(--accent); }
  .chat-send { background: var(--accent); color: #000; border: none; border-radius: 8px;
               padding: 10px 16px; font-weight: 600; cursor: pointer; font-size: 14px; white-space: nowrap; }
  .chat-send:disabled { opacity: 0.4; cursor: not-allowed; }
  .chat-save-insight { background: none; border: 1px solid var(--border); border-radius: 6px;
                       padding: 3px 8px; font-size: 11px; color: var(--muted); cursor: pointer;
                       margin-top: 4px; transition: all 0.15s; }
  .chat-save-insight:hover { border-color: var(--green); color: var(--green); }
  .chat-save-insight.saved { border-color: var(--green); color: var(--green); pointer-events: none; }
  .chat-feedback-row { display: flex; gap: 6px; margin-top: 4px; align-items: center; }
  .chat-fb-btn { background: none; border: 1px solid var(--border); border-radius: 6px;
                 padding: 2px 8px; font-size: 12px; cursor: pointer; transition: all 0.15s; color: var(--muted); }
  .chat-fb-btn:hover { border-color: var(--accent); }
  .chat-fb-btn.fb-good { border-color: var(--green); color: var(--green); }
  .chat-fb-btn.fb-bad { border-color: var(--red); color: var(--red); }
  .chat-fb-btn:disabled { opacity: 0.4; cursor: not-allowed; }
  .briefing-content h2, .briefing-content h3 { cursor: pointer; position: relative; }
  .briefing-content h2:hover, .briefing-content h3:hover { color: var(--accent); }
  .briefing-content h2::after, .briefing-content h3::after {
    content: '💬'; font-size: 12px; margin-left: 8px; opacity: 0; transition: opacity 0.15s; }
  .briefing-content h2:hover::after, .briefing-content h3:hover::after { opacity: 0.6; }
</style>
</head>
<body>
<div class="app">

<header>
  <div class="logo">AI Macro Trader <span id="header-user"></span></div>
  <div class="header-right desktop-nav">
    <button class="btn btn-secondary btn-sm" onclick="showSection('history')">History</button>
    <button class="btn btn-secondary btn-sm" onclick="showSection('settings')">Settings</button>
    <a href="/logout" class="btn btn-secondary btn-sm">Sign out</a>
  </div>
</header>

<main>

  <!-- HOME -->
  <div id="sec-home" class="section active">
    <div class="card">
      <div class="card-title">Today's Status</div>
      <div class="row" id="status-row"><div class="spinner"></div></div>
    </div>
    <div class="card admin-only">
      <div class="card-title">Generate Briefing</div>
      <div class="generate-area">
        <button id="gen-btn" class="btn btn-primary btn-large" onclick="generateBriefing()">
          Generate Today's Briefing
        </button>
        <div class="sub" id="gen-sub">Searches live macro news + streams directly to you</div>
      </div>
      <div id="log-box" class="log-box"></div>
    </div>
    <div id="briefing-card" class="card" style="display:none">
      <div class="card-title">
        <div class="row">
          <span id="briefing-card-title">Today's Briefing</span>
          <div class="spacer"></div>
        </div>
      </div>
      <div id="briefing-content" class="briefing-content"></div>
    </div>
    <div id="feedback-card" class="feedback-card admin-only" style="display:none">
      <div class="feedback-card-title">
        <svg width="14" height="14" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2"><path d="M21 15a2 2 0 01-2 2H7l-4 4V5a2 2 0 012-2h14a2 2 0 012 2z"/></svg>
        Trade Feedback — help improve future briefings
      </div>
      <p style="font-size:12px;color:var(--muted);margin-bottom:14px;line-height:1.5">
        Rate each trade idea. Thumbs down = wrong direction or poorly structured. Your notes teach future briefings.
      </p>
      <div id="feedback-items"></div>
      <div style="display:flex;align-items:center;gap:12px;margin-top:6px">
        <button class="feedback-save-btn" onclick="saveFeedback()">Save Feedback</button>
        <span class="feedback-saved-msg" id="feedback-saved-msg">Saved — next briefing will learn from this</span>
      </div>
    </div>
  </div>

  <!-- HISTORY -->
  <div id="sec-history" class="section">
    <div class="card">
      <div class="card-title">Past Briefings</div>
      <div id="history-list"><div class="empty">No briefings yet.</div></div>
    </div>
  </div>

  <!-- SETTINGS -->
  <div id="sec-settings" class="section">

    <!-- Schedule (admin only) -->
    <div class="card admin-only" id="card-schedule">
      <div class="card-title">Schedule</div>
      <div class="row" id="schedule-row" style="margin-bottom:14px"><div class="spinner"></div></div>
      <div class="row" style="gap:10px">
        <button class="btn btn-secondary" onclick="pauseBriefings()">Pause</button>
        <button class="btn btn-primary" onclick="resumeBriefings()">Resume</button>
      </div>
    </div>

    <!-- Knowledge Base (admin only) -->
    <div class="card admin-only" id="card-kb">
      <div class="card-title">Knowledge Base</div>
      <p style="font-size:12px;color:var(--muted);line-height:1.6;margin-bottom:14px">
        Upload research documents (.pdf, .docx, .txt). Each is summarized <em>once</em> using a lightweight AI call,
        then those notes are injected into every future briefing. Toggle off to exclude without deleting.
      </p>
      <div id="doc-list"><div class="empty" style="padding:12px 0">No documents uploaded yet.</div></div>
      <div class="upload-area">
        <div class="row" style="gap:8px;align-items:center;flex-wrap:wrap">
          <select id="upload-doc-type" class="doc-type-select dt-guide" onchange="this.className='doc-type-select dt-'+this.value">
            <option value="guide">Market Guide</option>
            <option value="tactical">Tactical Context</option>
            <option value="reference">Reference</option>
          </select>
          <input type="file" id="doc-upload-input" accept=".pdf,.docx,.txt" style="display:none" onchange="uploadDocument(this)">
          <button class="btn btn-primary btn-sm" onclick="document.getElementById('doc-upload-input').click()" id="upload-btn">
            + Upload Document
          </button>
        </div>
        <div style="font-size:11px;color:var(--muted);margin-top:6px;line-height:1.5">
          <strong>Tactical Context</strong> — headline news, speeches, short-term views<br>
          <strong>Market Guide</strong> — outlooks, structural themes, longer-term macro views<br>
          <strong>Reference</strong> — research papers, definitions, informational context
        </div>
        <div class="upload-status" id="upload-status"></div>
      </div>
    </div>

    <!-- Saved Insights (admin only) -->
    <div class="card admin-only" id="card-insights">
      <div class="card-title">Saved Insights</div>
      <p style="font-size:12px;color:var(--muted);line-height:1.6;margin-bottom:14px">
        Lessons saved from chat conversations. These are injected into every future briefing as permanent knowledge.
      </p>
      <div id="insights-list"><div class="empty" style="padding:12px 0">No insights saved yet.</div></div>
    </div>

    <!-- User Management (admin only) -->
    <div class="card admin-only" id="card-users">
      <div class="card-title">User Access</div>
      <p style="font-size:12px;color:var(--muted);line-height:1.6;margin-bottom:14px">
        Generate single-use invite links to grant access. Users create their own account with the link.
        You can disable or remove accounts at any time.
      </p>

      <div style="margin-bottom:20px">
        <div style="font-size:11px;font-weight:600;letter-spacing:0.06em;text-transform:uppercase;color:var(--muted);margin-bottom:10px">Create Invite</div>
        <div class="row" style="gap:8px">
          <input type="text" class="invite-note-input" id="invite-note" placeholder="Note (e.g. John Smith, Goldman)">
          <button class="btn btn-primary btn-sm" onclick="createInvite()">Generate Link</button>
        </div>
        <div class="status-msg" id="invite-status"></div>
      </div>

      <div style="margin-bottom:20px">
        <div style="font-size:11px;font-weight:600;letter-spacing:0.06em;text-transform:uppercase;color:var(--muted);margin-bottom:10px">Pending Invites</div>
        <div id="invites-list"><div class="empty" style="padding:8px 0">No pending invites.</div></div>
      </div>

      <div>
        <div style="font-size:11px;font-weight:600;letter-spacing:0.06em;text-transform:uppercase;color:var(--muted);margin-bottom:10px">Active Users</div>
        <div id="users-list"><div class="empty" style="padding:8px 0">No users yet.</div></div>
      </div>
    </div>

    <!-- Storage (admin only) -->
    <div class="card admin-only" id="card-storage">
      <div class="card-title">Persistent Storage (Railway Volume)</div>
      <p style="font-size:12px;color:var(--muted);line-height:1.6;margin-bottom:10px">
        By default, Railway resets the <code style="background:#1a1a1a;padding:1px 5px;border-radius:4px;font-family:monospace;font-size:11px;color:var(--accent)">data/</code> folder on every code deploy.
        Set up a Volume once to make everything permanent:
      </p>
      <div class="storage-steps">
        <strong>One-time setup:</strong><br>
        1. Open your <strong>Railway project dashboard</strong><br>
        2. Click your service → <strong>Volumes</strong> tab → <strong>Add Volume</strong><br>
        3. Set mount path: <code>/app/data</code><br>
        4. Click <strong>Deploy</strong><br><br>
        After this, all uploads, briefings, users &amp; feedback survive code deploys forever.
      </div>
    </div>

    <!-- About (everyone) -->
    <div class="card">
      <div class="card-title">About</div>
      <p style="font-size:13px;color:var(--muted);line-height:1.6">
        <strong style="color:var(--text)">AI Macro Trader</strong> — Daily briefings covering rates, FX, and cross-currency basis.<br><br>
        Briefings auto-generate at <strong style="color:var(--text)">6 AM ET Mon–Fri</strong> using the Anthropic API + live web search.
        Two learning systems run in parallel: section/trade feedback and uploaded knowledge base documents are both injected into every prompt.
      </p>
    </div>

  </div>

</main>

<nav class="tab-bar">
  <div class="tab-bar-inner">
    <button class="tab-btn active" id="tab-home" onclick="showSection('home')">
      <svg viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="1.8"><path d="M3 9.5L12 3l9 6.5V20a1 1 0 01-1 1H4a1 1 0 01-1-1V9.5z"/></svg>
      Home
    </button>
    <button class="tab-btn" id="tab-history" onclick="showSection('history')">
      <svg viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="1.8"><rect x="3" y="4" width="18" height="18" rx="2"/><path d="M16 2v4M8 2v4M3 10h18"/></svg>
      History
    </button>
    <button class="tab-btn admin-only" id="tab-settings" onclick="showSection('settings')">
      <svg viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="1.8"><circle cx="12" cy="12" r="3"/><path d="M12 2v2m0 16v2M4.22 4.22l1.42 1.42m12.72 12.72l1.42 1.42M2 12h2m16 0h2M4.22 19.78l1.42-1.42M18.36 5.64l1.42-1.42"/></svg>
      Settings
    </button>
  </div>
</nav>

<!-- CHAT PANEL -->
<div class="chat-overlay" id="chat-overlay" onclick="closeChat()"></div>
<div class="chat-panel" id="chat-panel">
  <div class="chat-header">
    <span class="chat-header-title">💬 Macro LLM — Ask about this briefing</span>
    <button class="chat-close" onclick="closeChat()">✕</button>
  </div>
  <div class="chat-context" id="chat-context"></div>
  <div class="chat-messages" id="chat-messages"></div>
  <div class="chat-input-row">
    <textarea class="chat-input" id="chat-input" rows="1"
      placeholder="Ask a question, challenge a trade idea, dig deeper..."
      onkeydown="if(event.key==='Enter'&&!event.shiftKey){event.preventDefault();sendChat()}"></textarea>
    <button class="chat-send" id="chat-send-btn" onclick="sendChat()">Send</button>
  </div>
</div>

</div>
<script>
let currentSection = 'home';
let currentBriefingFile = null;
let statusData = null;
let isAdmin = false;
let chatSectionContext = '';

async function initApp() {
  const r = await fetch('/api/me');
  if (r.ok) {
    const me = await r.json();
    isAdmin = me.is_admin;
    const el = document.getElementById('header-user');
    if (el) el.textContent = me.username ? `// ${me.username}` : '';
    if (isAdmin) {
      document.querySelectorAll('.admin-only').forEach(el => el.classList.add('visible'));
    }
  }
  loadStatus();
}

function showSection(name) {
  document.querySelectorAll('.section').forEach(s => s.classList.remove('active'));
  document.querySelectorAll('.tab-btn').forEach(b => b.classList.remove('active'));
  document.getElementById('sec-' + name).classList.add('active');
  const tab = document.getElementById('tab-' + name);
  if (tab) tab.classList.add('active');
  currentSection = name;
  if (name === 'history') loadHistory();
  if (name === 'settings') loadSettings();
}

async function loadStatus() {
  const row = document.getElementById('status-row');
  try {
    const r = await fetch('/api/status');
    if (!r.ok) throw new Error(`${r.status}`);
    statusData = await r.json();
    let pills = '';
    pills += statusData.paused
      ? `<span class="pill pill-red"><span class="dot"></span>Paused</span>`
      : `<span class="pill pill-green"><span class="dot"></span>Active</span>`;
    pills += statusData.today_done
      ? `<span class="pill pill-gold"><span class="dot"></span>Today's briefing ready</span>`
      : `<span class="pill pill-red"><span class="dot"></span>Not yet generated today</span>`;
    row.innerHTML = pills;
    if (statusData.today_done) {
      const fname = `macro-briefing-${statusData.today_date}.md`;
      await loadBriefing(fname);
      document.getElementById('gen-btn').textContent = "Regenerate Today's Briefing";
      document.getElementById('gen-sub').textContent = `Briefing exists for ${statusData.today_date}. Click to regenerate.`;
    }
  } catch(e) {
    row.innerHTML = `<span style="color:var(--red);font-size:13px">Error: ${e.message}</span>`;
  }
}

async function generateBriefing() {
  const btn = document.getElementById('gen-btn');
  const logBox = document.getElementById('log-box');
  const sub = document.getElementById('gen-sub');
  if (statusData && statusData.today_done) {
    await fetch('/api/delete-today', { method: 'POST' });
  }
  btn.disabled = true;
  btn.innerHTML = '<span class="spinner"></span> Generating...';
  sub.textContent = 'Searching live macro news — takes 1-3 minutes...';
  logBox.innerHTML = '';
  logBox.classList.add('visible');
  const evtSource = new EventSource('/api/generate');
  evtSource.onmessage = (e) => {
    const msg = e.data;
    if (msg.startsWith('__DONE_OK__')) {
      evtSource.close();
      const fname = msg.replace('__DONE_OK__', '').trim();
      btn.disabled = false;
      btn.innerHTML = "Generate Today's Briefing";
      sub.textContent = 'Done!';
      appendLog(logBox, 'Complete.', 'ok');
      loadBriefing(fname);
      loadStatus();
      return;
    }
    if (msg === '__DONE_ERROR__') {
      evtSource.close();
      btn.disabled = false;
      btn.innerHTML = "Generate Today's Briefing";
      sub.textContent = 'Generation failed. Check log.';
      appendLog(logBox, 'Failed.', 'err');
      return;
    }
    appendLog(logBox, msg, msg.startsWith('Error') ? 'err' : '');
  };
  evtSource.onerror = () => {
    evtSource.close();
    btn.disabled = false;
    btn.innerHTML = "Generate Today's Briefing";
    appendLog(logBox, 'Connection lost.', 'err');
  };
}

function appendLog(box, text, type) {
  const line = document.createElement('div');
  if (type === 'ok') line.className = 'log-line-ok';
  if (type === 'err') line.className = 'log-line-err';
  line.textContent = text;
  box.appendChild(line);
  box.scrollTop = box.scrollHeight;
}

async function loadBriefing(filename) {
  currentBriefingFile = filename;
  const briefingResp = await fetch(`/api/briefing/${filename}`);
  if (!briefingResp.ok) return;
  const data = await briefingResp.json();
  let allFeedback = [];
  let sectionRatings = {};
  if (isAdmin) {
    const feedbackResp = await fetch(`/api/feedback?filename=${filename}`);
    allFeedback = feedbackResp.ok ? await feedbackResp.json() : [];
    for (const entry of allFeedback) {
      if (entry.section) {
        sectionRatings[entry.section] = { rating: entry.rating, note: entry.note || '' };
      }
    }
  }
  const dateStr = filename.replace('macro-briefing-', '').replace('.md', '');
  document.getElementById('briefing-card-title').textContent = `Briefing — ${dateStr}`;
  document.getElementById('briefing-content').innerHTML = renderMarkdown(data.content, sectionRatings);
  document.getElementById('briefing-card').style.display = 'block';
  if (isAdmin) await renderFeedback(filename, data.content, allFeedback);
}

function renderMarkdown(md, sectionRatings) {
  sectionRatings = sectionRatings || {};
  let html = md
    .replace(/&/g, '&amp;').replace(/</g, '&lt;').replace(/>/g, '&gt;')
    .replace(/^# (.+)$/gm, '<h1>$1</h1>')
    .replace(/^## (.+)$/gm, (match, title) => {
      const key = title.trim();
      if (!isAdmin) return `<h2>${key}</h2>`;
      const saved = sectionRatings[key] || {};
      const upCls   = saved.rating === 'up'   ? 'sf-up'   : '';
      const downCls = saved.rating === 'down' ? 'sf-down' : '';
      const noteVal = (saved.note || '').replace(/"/g, '&quot;');
      return `<h2>${key}</h2><div class="section-feedback-row" data-section="${key.replace(/"/g,'&quot;')}">` +
        `<span class="sf-label">Section:</span>` +
        `<button class="sf-thumb ${upCls}" data-dir="up" onclick="sfThumb(this)">👍</button>` +
        `<button class="sf-thumb ${downCls}" data-dir="down" onclick="sfThumb(this)">👎</button>` +
        `<textarea class="sf-note" placeholder="Notes on this section...">${noteVal}</textarea>` +
        `<button class="sf-save" onclick="sfSave(this)">Save</button>` +
        `<span class="sf-saved">Saved</span>` +
        `</div>`;
    })
    .replace(/^### (.+)$/gm, '<h3>$1</h3>')
    .replace(/\*\*(.+?)\*\*/g, '<strong>$1</strong>')
    .replace(/\*(.+?)\*/g, '<em>$1</em>')
    .replace(/\n{2,}/g, '</p><p>')
    .replace(/^/, '<p>').replace(/$/, '</p>');
  return html;
}

async function loadHistory() {
  const r = await fetch('/api/briefings');
  const files = await r.json();
  const list = document.getElementById('history-list');
  if (!files.length) { list.innerHTML = '<div class="empty">No briefings yet.</div>'; return; }
  list.innerHTML = files.map(f => {
    const d = f.replace('macro-briefing-', '').replace('.md', '');
    return `<div class="history-item">
      <span class="history-date">${d}</span>
      <div class="history-actions">
        <button class="btn btn-secondary btn-sm" onclick="viewHistory('${f}')">View</button>
      </div>
    </div>`;
  }).join('');
}

async function viewHistory(fname) {
  await loadBriefing(fname);
  showSection('home');
  window.scrollTo(0, document.getElementById('briefing-card').offsetTop - 80);
}

async function loadSettings() {
  const r = await fetch('/api/status');
  const s = await r.json();
  const schedRow = document.getElementById('schedule-row');
  if (schedRow) {
    schedRow.innerHTML = s.paused
      ? `<span class="pill pill-red"><span class="dot"></span>Paused</span>`
      : `<span class="pill pill-green"><span class="dot"></span>Active — Mon-Fri 6 AM ET</span>`;
  }
  if (isAdmin) {
    await loadDocuments();
    await loadInsightsList();
    await loadUsersAndInvites();
  }
}

async function loadInsightsList() {
  const r = await fetch('/api/insights');
  if (!r.ok) return;
  const insights = await r.json();
  const list = document.getElementById('insights-list');
  if (!list) return;
  if (!insights.length) {
    list.innerHTML = '<div class="empty" style="padding:12px 0">No insights saved yet. Use the 💬 chat to discuss briefing sections, then save key takeaways.</div>';
    return;
  }
  list.innerHTML = insights.map((ins, i) => `
    <div class="doc-item">
      <span class="doc-name" style="font-size:12px;line-height:1.5">${escapeHtmlBasic(ins.insight)}</span>
      <span class="doc-date">${ins.date || ins.saved_at || ''}</span>
      <button class="doc-delete" onclick="deleteInsight(${i})">✕</button>
    </div>`).join('');
}

function escapeHtmlBasic(s) {
  return s.replace(/&/g,'&amp;').replace(/</g,'&lt;').replace(/>/g,'&gt;');
}

async function deleteInsight(idx) {
  if (!confirm('Delete this insight? It will no longer be injected into future briefings.')) return;
  const r = await fetch(`/api/insights/${idx}/delete`, { method: 'POST' });
  if (r.ok) await loadInsightsList();
}

async function pauseBriefings() {
  await fetch('/api/pause', { method: 'POST', headers: {'Content-Type':'application/json'}, body: '{}' });
  loadStatus(); loadSettings();
}

async function resumeBriefings() {
  await fetch('/api/resume', { method: 'POST', headers: {'Content-Type':'application/json'}, body: '{}' });
  loadStatus(); loadSettings();
}

// ── Document upload ──
async function loadDocuments() {
  const r = await fetch('/api/documents');
  if (!r.ok) return;
  const docs = await r.json();
  const list = document.getElementById('doc-list');
  if (!list) return;
  if (!docs.length) {
    list.innerHTML = '<div class="empty" style="padding:12px 0">No documents uploaded yet.</div>';
    return;
  }
  list.innerHTML = docs.map(d => `
    <div class="doc-item" id="doc-${d.id}">
      <span class="doc-name">${d.title}</span>
      <span class="doc-date">${d.uploaded_at}</span>
      <div class="doc-actions">
        <select class="doc-type-select dt-${d.doc_type || 'guide'}"
          onchange="setDocType('${d.id}', this.value, this)">
          <option value="guide"${d.doc_type === 'guide' || !d.doc_type ? ' selected' : ''}>Market Guide</option>
          <option value="tactical"${d.doc_type === 'tactical' ? ' selected' : ''}>Tactical Context</option>
          <option value="reference"${d.doc_type === 'reference' ? ' selected' : ''}>Reference</option>
        </select>
        <button class="toggle-pill ${d.active ? 'active' : 'inactive'}"
          onclick="toggleDocument('${d.id}', this)">${d.active ? 'Active' : 'Off'}</button>
        <button class="doc-delete" onclick="deleteDocument('${d.id}')">✕</button>
      </div>
    </div>`).join('');
}

async function uploadDocument(input) {
  const file = input.files[0];
  if (!file) return;
  const status = document.getElementById('upload-status');
  const btn = document.getElementById('upload-btn');
  status.className = 'upload-status';
  status.textContent = 'Uploading and processing... (~10 seconds)';
  btn.disabled = true;
  input.value = '';
  const docType = document.getElementById('upload-doc-type').value;
  const formData = new FormData();
  formData.append('file', file);
  formData.append('doc_type', docType);
  try {
    const r = await fetch('/api/upload', { method: 'POST', body: formData });
    const data = await r.json();
    if (r.ok && data.ok) {
      status.className = 'upload-status ok';
      status.textContent = `Processed: "${data.doc.title}" — injected into all future briefings`;
      await loadDocuments();
    } else {
      status.className = 'upload-status err';
      status.textContent = `Error: ${data.error || 'Upload failed'}`;
    }
  } catch(e) {
    status.className = 'upload-status err';
    status.textContent = `Error: ${e.message}`;
  }
  btn.disabled = false;
  setTimeout(() => { status.textContent = ''; status.className = 'upload-status'; }, 8000);
}

async function toggleDocument(docId, btn) {
  const r = await fetch(`/api/documents/${docId}/toggle`, { method: 'POST' });
  if (!r.ok) return;
  const data = await r.json();
  btn.className = `toggle-pill ${data.active ? 'active' : 'inactive'}`;
  btn.textContent = data.active ? 'Active' : 'Off';
}

async function deleteDocument(docId) {
  if (!confirm('Remove this document from the knowledge base?')) return;
  const r = await fetch(`/api/documents/${docId}/delete`, { method: 'POST' });
  if (r.ok) { await loadDocuments(); }
}

async function setDocType(docId, newType, sel) {
  const r = await fetch(`/api/documents/${docId}/type`, {
    method: 'POST',
    headers: {'Content-Type': 'application/json'},
    body: JSON.stringify({doc_type: newType})
  });
  if (!r.ok) return;
  const data = await r.json();
  sel.className = `doc-type-select dt-${data.doc_type}`;
}

// ── User management ──
async function loadUsersAndInvites() {
  await Promise.all([loadInvites(), loadUsers()]);
}

async function loadInvites() {
  const r = await fetch('/api/admin/invites');
  if (!r.ok) return;
  const invites = await r.json();
  const list = document.getElementById('invites-list');
  if (!list) return;
  const pending = invites.filter(i => !i.used);
  if (!pending.length) {
    list.innerHTML = '<div class="empty" style="padding:8px 0">No pending invites.</div>';
    return;
  }
  list.innerHTML = pending.map(inv => `
    <div class="invite-item" id="inv-${inv.token}">
      <div style="flex:1;min-width:0">
        <div class="invite-token">${inv.token}</div>
        <div class="invite-meta">${inv.note ? inv.note + ' · ' : ''}Created ${inv.created_at}</div>
        <div class="invite-link" id="inv-link-${inv.token}">${window.location.origin}/register?token=${inv.token}</div>
      </div>
      <div style="display:flex;gap:6px;flex-shrink:0">
        <button class="copy-btn" id="copy-${inv.token}" onclick="copyInviteLink('${inv.token}')">Copy</button>
        <button class="doc-delete" onclick="deleteInvite('${inv.token}')">✕</button>
      </div>
    </div>`).join('');
}

async function loadUsers() {
  const r = await fetch('/api/admin/users');
  if (!r.ok) return;
  const users = await r.json();
  const list = document.getElementById('users-list');
  if (!list) return;
  if (!users.length) {
    list.innerHTML = '<div class="empty" style="padding:8px 0">No users yet.</div>';
    return;
  }
  list.innerHTML = users.map(u => `
    <div class="user-item" id="user-${u.username}">
      <div style="flex:1">
        <span class="user-name">${u.username}</span>
        ${u.note ? `<span class="user-note"> — ${u.note}</span>` : ''}
        <div class="user-date">Joined ${u.created_at}</div>
      </div>
      <div style="display:flex;gap:6px;align-items:center">
        <button class="toggle-pill ${u.active ? 'active' : 'inactive'}"
          onclick="toggleUser('${u.username}', this)">${u.active ? 'Active' : 'Disabled'}</button>
        <button class="doc-delete" onclick="deleteUser('${u.username}')">✕</button>
      </div>
    </div>`).join('');
}

async function createInvite() {
  const note = document.getElementById('invite-note').value.trim();
  const status = document.getElementById('invite-status');
  const r = await fetch('/api/admin/invite', {
    method: 'POST',
    headers: {'Content-Type':'application/json'},
    body: JSON.stringify({ note })
  });
  if (r.ok) {
    const data = await r.json();
    document.getElementById('invite-note').value = '';
    status.className = 'status-msg ok';
    status.textContent = `Invite created: ${data.token}`;
    setTimeout(() => { status.textContent = ''; }, 5000);
    await loadInvites();
  } else {
    status.className = 'status-msg err';
    status.textContent = 'Failed to create invite.';
  }
}

function copyInviteLink(token) {
  const link = `${window.location.origin}/register?token=${token}`;
  navigator.clipboard.writeText(link).then(() => {
    const btn = document.getElementById(`copy-${token}`);
    if (btn) { btn.textContent = 'Copied!'; btn.classList.add('copied'); }
    setTimeout(() => {
      if (btn) { btn.textContent = 'Copy'; btn.classList.remove('copied'); }
    }, 2500);
  });
}

async function deleteInvite(token) {
  if (!confirm('Delete this invite link?')) return;
  const r = await fetch(`/api/admin/invite/${token}/delete`, { method: 'POST' });
  if (r.ok) await loadInvites();
}

async function toggleUser(username, btn) {
  const r = await fetch(`/api/admin/users/${username}/toggle`, { method: 'POST' });
  if (!r.ok) return;
  const data = await r.json();
  btn.className = `toggle-pill ${data.active ? 'active' : 'inactive'}`;
  btn.textContent = data.active ? 'Active' : 'Disabled';
}

async function deleteUser(username) {
  if (!confirm(`Remove access for ${username}? They will no longer be able to sign in.`)) return;
  const r = await fetch(`/api/admin/users/${username}/delete`, { method: 'POST' });
  if (r.ok) await loadUsers();
}

// ── Feedback ──
function extractTrades(md) {
  const section = md.match(/##\s*Trade Construction Context([\s\S]*?)(?=\n##\s|\n#\s|$)/i);
  if (!section) return [];
  const text = section[1].trim();
  if (!text) return [];
  const raw = text.split(/\n(?=\d+\.|[-*]|\*\*|Given\s)/i)
    .map(s => s.replace(/^[\d.\-*]+\s*/, '').trim())
    .filter(s => s.length > 30);
  return raw.length ? raw : [text];
}

async function renderFeedback(filename, md, existingFeedback) {
  const trades = extractTrades(md);
  const card = document.getElementById('feedback-card');
  const container = document.getElementById('feedback-items');
  if (!trades.length) { card.style.display = 'none'; return; }
  const existing = existingFeedback || [];
  const tradeEntries = existing.filter(e => !e.section);
  container.innerHTML = '';
  trades.forEach((trade, i) => {
    const saved = tradeEntries[i] || {};
    const upActive   = saved.rating === 'up'   ? 'active-up'   : '';
    const downActive = saved.rating === 'down' ? 'active-down' : '';
    const item = document.createElement('div');
    item.className = 'trade-feedback-item';
    item.dataset.index = i;
    item.dataset.trade = trade;
    item.innerHTML = `
      <div class="trade-text">${trade.replace(/</g,'&lt;').replace(/>/g,'&gt;')}</div>
      <div class="feedback-controls">
        <button class="thumb-btn ${upActive}"   data-dir="up"   onclick="toggleThumb(this)">👍</button>
        <button class="thumb-btn ${downActive}" data-dir="down" onclick="toggleThumb(this)">👎</button>
        <textarea class="feedback-note" placeholder="Why? e.g. wrong tenor, direction off, crowded...">${saved.note||''}</textarea>
      </div>`;
    container.appendChild(item);
  });
  card.style.display = 'block';
}

function sfThumb(btn) {
  const row = btn.closest('.section-feedback-row');
  const dir = btn.dataset.dir;
  const wasActive = btn.classList.contains(`sf-${dir}`);
  row.querySelectorAll('.sf-thumb').forEach(b => b.classList.remove('sf-up','sf-down'));
  if (!wasActive) btn.classList.add(`sf-${dir}`);
}

async function sfSave(btn) {
  if (!currentBriefingFile) return;
  const row = btn.closest('.section-feedback-row');
  const section = row.dataset.section;
  const upBtn   = row.querySelector('[data-dir="up"]');
  const downBtn = row.querySelector('[data-dir="down"]');
  const note    = row.querySelector('.sf-note').value.trim();
  const savedSpan = row.querySelector('.sf-saved');
  let rating = null;
  if (upBtn.classList.contains('sf-up'))     rating = 'up';
  if (downBtn.classList.contains('sf-down')) rating = 'down';
  const r = await fetch('/api/feedback', {
    method: 'POST',
    headers: {'Content-Type':'application/json'},
    body: JSON.stringify({ filename: currentBriefingFile, entries: [{ section, rating, note }] })
  });
  if (r.ok) {
    savedSpan.style.display = 'inline';
    setTimeout(() => { savedSpan.style.display = 'none'; }, 3000);
  }
}

function toggleThumb(btn) {
  const item = btn.closest('.trade-feedback-item');
  const dir = btn.dataset.dir;
  const wasActive = btn.classList.contains(`active-${dir}`);
  item.querySelectorAll('.thumb-btn').forEach(b => b.classList.remove('active-up','active-down'));
  if (!wasActive) btn.classList.add(`active-${dir}`);
}

async function saveFeedback() {
  if (!currentBriefingFile) return;
  const saveBtn = document.querySelector('.feedback-save-btn');
  const savedMsg = document.getElementById('feedback-saved-msg');
  saveBtn.disabled = true;
  const entries = [];
  document.querySelectorAll('.trade-feedback-item').forEach(item => {
    const upBtn   = item.querySelector('[data-dir="up"]');
    const downBtn = item.querySelector('[data-dir="down"]');
    const note    = item.querySelector('.feedback-note').value.trim();
    let rating = null;
    if (upBtn.classList.contains('active-up'))     rating = 'up';
    if (downBtn.classList.contains('active-down')) rating = 'down';
    entries.push({ trade: item.dataset.trade, rating, note });
  });
  const r = await fetch('/api/feedback', {
    method: 'POST',
    headers: {'Content-Type':'application/json'},
    body: JSON.stringify({ filename: currentBriefingFile, entries })
  });
  saveBtn.disabled = false;
  if (r.ok) {
    savedMsg.style.display = 'inline';
    setTimeout(() => { savedMsg.style.display = 'none'; }, 4000);
  }
}

// ── Chat panel ──
function openChat(sectionTitle, sectionContent) {
  if (!isAdmin) return;
  chatSectionContext = sectionContent || '';
  const ctxEl = document.getElementById('chat-context');
  if (sectionTitle) {
    ctxEl.textContent = `Discussing: ${sectionTitle}`;
    ctxEl.style.display = 'block';
  } else {
    ctxEl.style.display = 'none';
  }
  document.getElementById('chat-overlay').classList.add('open');
  document.getElementById('chat-panel').classList.add('open');
  document.getElementById('chat-input').focus();
  loadChatHistory();
}

function closeChat() {
  document.getElementById('chat-overlay').classList.remove('open');
  document.getElementById('chat-panel').classList.remove('open');
}

async function loadChatHistory() {
  if (!currentBriefingFile) return;
  const dateStr = currentBriefingFile.replace('macro-briefing-', '').replace('.md', '');
  const r = await fetch(`/api/chat/history/${dateStr}`);
  if (!r.ok) return;
  const messages = await r.json();
  const container = document.getElementById('chat-messages');
  container.innerHTML = '';
  for (const msg of messages) {
    appendChatMessage(msg.role, msg.content, msg.ts, true);
  }
  container.scrollTop = container.scrollHeight;
}

function appendChatMessage(role, content, ts, skipFeedback) {
  const container = document.getElementById('chat-messages');
  const div = document.createElement('div');
  div.className = `chat-msg chat-msg-${role}`;
  let html = `<div class="chat-bubble">${escapeHtml(content)}</div>`;
  if (ts) html += `<div class="chat-msg-time">${ts}</div>`;
  if (role === 'assistant') {
    const insightText = content.replace(/\n/g, ' ').substring(0, 300);
    html += `<div class="chat-feedback-row">`;
    if (!skipFeedback) {
      html += `<button class="chat-fb-btn" onclick="chatFeedback(this,'good')" title="Good response">👍</button>`;
      html += `<button class="chat-fb-btn" onclick="chatFeedback(this,'bad')" title="Bad response">👎</button>`;
    }
    html += `<button class="chat-save-insight" onclick="saveInsight(this, \`${insightText.replace(/`/g,'\\`').replace(/\\/g,'\\\\')}\`)">💡 Save as insight</button>`;
    html += `</div>`;
  }
  div.innerHTML = html;
  container.appendChild(div);
  container.scrollTop = container.scrollHeight;
}

async function chatFeedback(btn, feedback) {
  const row = btn.closest('.chat-feedback-row');
  row.querySelectorAll('.chat-fb-btn').forEach(b => { b.disabled = true; b.classList.remove('fb-good','fb-bad'); });
  btn.classList.add(feedback === 'good' ? 'fb-good' : 'fb-bad');
  try {
    await fetch('/api/chat/feedback', {
      method: 'POST',
      headers: {'Content-Type': 'application/json'},
      body: JSON.stringify({ feedback })
    });
  } catch(e) { console.error('Feedback error:', e); }
}

function escapeHtml(text) {
  return text.replace(/&/g,'&amp;').replace(/</g,'&lt;').replace(/>/g,'&gt;')
    .replace(/\*\*(.+?)\*\*/g, '<strong>$1</strong>')
    .replace(/\n/g, '<br>');
}

async function sendChat() {
  const input = document.getElementById('chat-input');
  const btn = document.getElementById('chat-send-btn');
  const msg = input.value.trim();
  if (!msg || !currentBriefingFile) return;
  const dateStr = currentBriefingFile.replace('macro-briefing-', '').replace('.md', '');

  input.value = '';
  btn.disabled = true;
  appendChatMessage('user', msg, new Date().toLocaleTimeString([], {hour:'2-digit',minute:'2-digit'}));

  // Add typing indicator
  const container = document.getElementById('chat-messages');
  const typingDiv = document.createElement('div');
  typingDiv.className = 'chat-msg chat-msg-assistant';
  typingDiv.id = 'chat-typing';
  typingDiv.innerHTML = '<div class="chat-typing">Macro LLM reasoning...</div>';
  container.appendChild(typingDiv);
  container.scrollTop = container.scrollHeight;

  try {
    const r = await fetch('/api/chat', {
      method: 'POST',
      headers: {'Content-Type': 'application/json'},
      body: JSON.stringify({
        briefing_date: dateStr,
        message: msg,
        section_context: chatSectionContext
      })
    });

    const reader = r.body.getReader();
    const decoder = new TextDecoder();
    let fullResponse = '';
    let buffer = '';

    // Replace typing indicator with streaming bubble
    typingDiv.innerHTML = '<div class="chat-bubble" id="chat-stream-bubble"></div>';
    const bubble = document.getElementById('chat-stream-bubble');

    while (true) {
      const {done, value} = await reader.read();
      if (done) break;
      buffer += decoder.decode(value, {stream: true});
      const lines = buffer.split('\n');
      buffer = lines.pop();
      for (const line of lines) {
        if (!line.startsWith('data: ')) continue;
        try {
          const data = JSON.parse(line.slice(6));
          if (data.chunk) {
            fullResponse += data.chunk;
            bubble.innerHTML = escapeHtml(fullResponse);
            container.scrollTop = container.scrollHeight;
          }
          if (data.done) {
            // Replace streaming div with proper message
            typingDiv.remove();
            appendChatMessage('assistant', fullResponse, new Date().toLocaleTimeString([], {hour:'2-digit',minute:'2-digit'}));
          }
          if (data.error) {
            typingDiv.innerHTML = `<div class="chat-bubble" style="color:var(--red)">Error: ${data.error}</div>`;
          }
        } catch(e) {}
      }
    }
  } catch(e) {
    const existing = document.getElementById('chat-typing');
    if (existing) existing.innerHTML = `<div class="chat-bubble" style="color:var(--red)">Error: ${e.message}</div>`;
  }
  btn.disabled = false;
  input.focus();
}

async function saveInsight(btn, text) {
  if (!currentBriefingFile || !text) return;
  btn.disabled = true;
  btn.textContent = 'Saving...';
  const dateStr = currentBriefingFile.replace('macro-briefing-', '').replace('.md', '');
  try {
    const r = await fetch('/api/chat/save-insight', {
      method: 'POST',
      headers: {'Content-Type': 'application/json'},
      body: JSON.stringify({ insight: text, briefing_date: dateStr })
    });
    if (r.ok) {
      btn.textContent = '✓ Saved as insight';
      btn.classList.add('saved');
    } else {
      btn.textContent = '✗ Failed to save';
      btn.disabled = false;
    }
  } catch(e) {
    btn.textContent = '✗ Error';
    btn.disabled = false;
  }
}

// Make briefing headings clickable to open chat
document.addEventListener('click', function(e) {
  if (!isAdmin) return;
  const heading = e.target.closest('.briefing-content h2, .briefing-content h3');
  if (!heading) return;
  // Don't trigger if clicking on section feedback controls
  if (e.target.closest('.section-feedback-row')) return;
  const title = heading.textContent.replace('💬','').trim();
  // Grab the content up to the next heading of same or higher level
  let content = '';
  let el = heading.nextElementSibling;
  const level = heading.tagName;
  while (el) {
    if (el.tagName === 'H1' || el.tagName === 'H2' || (level === 'H3' && el.tagName === 'H3')) break;
    if (el.classList && el.classList.contains('section-feedback-row')) { el = el.nextElementSibling; continue; }
    content += el.textContent + '\n';
    el = el.nextElementSibling;
  }
  openChat(title, `## ${title}\n${content.trim()}`);
});

initApp();
</script>
</body>
</html>"""

if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5050))
    app.run(host="0.0.0.0", port=port, debug=False)
